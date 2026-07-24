#!/usr/bin/env python3
"""
redact_pii.py
-------------
Reads a .docx, finds PII, replaces it with consistent fake values, and
writes a redacted .docx.

Usage:
    python3 redact_pii.py INPUT.docx OUTPUT.docx [--log detections.json]

Design (see README.md for the full writeup):

  1. STRUCTURAL PASS  - a handful of tables/paragraphs in a document like
     this have unambiguous column headers ("Name", "Address", "Contact
     Person", "E-mail and Telephone", ...). Where we can identify that
     structure we redact the *whole cell* according to its column type.
     This is high precision/recall for exactly the rows that matter most
     (the Board of Directors table, the registered-office block, etc.)
     and it also feeds discovered names into a NameRegistry so the same
     person's name is caught later wherever it appears in free-flowing
     prose too.

  2. FREE-TEXT PASS - every remaining paragraph and table cell is run
     through a set of regex/heuristic detectors (detectors.py) plus an
     exact-match sweep for every name/org the structural pass discovered.
     Overlapping matches are resolved (longer span wins) and then the
     text is rewritten right-to-left so earlier offsets stay valid.

No NER model (spaCy/Presidio) is used -- see README.md, this sandbox has
no outbound network access so model downloads aren't possible here. If
you have network access, swapping the free-text pass for a
Presidio/spaCy analyzer is the natural upgrade path.
"""

import argparse
import json
import re
import sys
from pathlib import Path

import docx
from docx.table import Table
from docx.text.paragraph import Paragraph

from detectors import (
    Span,
    EmailDetector, IPAddressDetector, SSNDetector, CreditCardDetector,
    PhoneNumberDetector, DateOfBirthDetector, AddressDetector,
    HonorificNameDetector, CompanySuffixDetector,
)
from fake_data import FakeIdentityProvider

# ---------------------------------------------------------------------------
# Organisations we deliberately do NOT redact, and why (see README "Choices"
# section). Matched case-insensitively as substrings.
# ---------------------------------------------------------------------------
ORG_ALLOWLIST = [
    # The issuer itself (incl. former names) -- it is the subject of the
    # document, named in the title/on every page; redacting it would not
    # protect anyone and would make the document meaningless.
    "KSH International", "Bhandary Metal Extrusion",
    # Statutory/regulatory/market-infrastructure bodies and named Acts --
    # generic public institutions, not a confidential business relationship
    # (same category the assignment itself carves out for "Order"/"Ticket"
    # numbers).
    "Securities and Exchange Board of India", "SEBI",
    "National Stock Exchange", "NSE", "BSE Limited", "BSE",
    "Reserve Bank of India", "RBI", "Registrar of Companies", "RoC",
    "Ministry of Corporate Affairs", "Central Depository Services",
    "CDSL", "National Securities Depository", "NSDL",
    "Competition Commission of India", "Companies Act", "SEBI Act",
    "Depositories Act", "Income Tax Act", "Stock Exchanges",
]

# ---------------------------------------------------------------------------
# Column-header -> PII type, for the "structural pass" over docx tables.
# Header text is matched case-insensitively, substring match.
# ---------------------------------------------------------------------------
COLUMN_TYPE_MAP = [
    ("contact person", "PERSON"),
    ("name of the shareholder", "PERSON_OR_ORG"),
    ("name of the promoter", "PERSON_OR_ORG"),
    ("name of promoter", "PERSON_OR_ORG"),
    ("name of shareholder", "PERSON_OR_ORG"),
    ("name of related party", "PERSON_OR_ORG"),
    ("name of entity", "ORG"),
    ("name", "PERSON_OR_ORG"),          # generic fallback, checked last
    ("address", "ADDRESS"),
    ("e-mail and telephone", "EMAIL_PHONE"),
    ("email", "EMAIL"),
    ("e-mail", "EMAIL"),
    ("telephone", "PHONE_TEXT"),
]
# Columns we explicitly recognise but choose NOT to redact, with reasons.
COLUMN_SKIP = {
    "din": "public regulatory identifier (Director Identification Number), "
           "not in the required PII list -- analogous to the assignment's "
           "own Order/Ticket-number carve-out",
    "designation": "a job title alone does not identify a person",
}


_TRAILING_STOPWORDS = {
    "website", "email", "e-mail", "tel", "telephone", "phone", "fax",
    "contact", "address", "and", "compliance", "officer", "secretary",
}
_NON_NAME_EXACT = {
    "total", "name", "address", "promoters", "promoter group",
}
_GENERIC_LABEL_RE = re.compile(
    r"\b(payment|remuneration|expense|compensation|personnel|total|sub-total|"
    r"registration|sebi|equity shares?|face value|shareholding)\b",
    re.IGNORECASE,
)


def clean_name_candidate(raw: str) -> str:
    """Strip footnote markers (*, ^, &), tabs/newlines, and any trailing
    words that are clearly not part of a person's name (leaked in from
    'Contact Person: X\\nWebsite: ...'-style prose the regex over-matched)."""
    s = re.sub(r"[\t\n]+", " ", raw)
    s = re.sub(r"[\*\^&]+", "", s)
    s = re.sub(r"\s+", " ", s).strip().strip(",")
    words = s.split(" ")
    while words and words[-1].lower() in _TRAILING_STOPWORDS:
        words.pop()
    return " ".join(words).strip()


class NameRegistry:
    """Names/orgs discovered via structural anchors, ready for a
    document-wide exact-match sweep during the free-text pass."""

    def __init__(self):
        self.person_names = set()
        self.org_names = set()

    def add_person(self, name: str):
        name = clean_name_candidate(name)
        if not name or len(name) <= 2 or name.lower() in _NON_NAME_EXACT or " " not in name:
            return
        if len(name.split()) > 5 or _GENERIC_LABEL_RE.search(name):
            return  # too long / looks like an accounting label, not a name
        self.person_names.add(name)

    def add_org(self, name: str):
        name = clean_name_candidate(name)
        if not name or len(name) <= 2 or name.lower() in _NON_NAME_EXACT:
            return
        if len(name.split()) > 6 or _GENERIC_LABEL_RE.search(name):
            return
        self.org_names.add(name)

    def all_terms(self):
        """Longest-first so multi-word names are matched before any
        shorter substring of themselves."""
        terms = [(n, "PERSON") for n in self.person_names]
        terms += [(n, "COMPANY") for n in self.org_names]
        terms.sort(key=lambda t: -len(t[0]))
        return terms


def is_allowlisted(name: str) -> bool:
    low = name.lower()
    return any(a.lower() in low for a in ORG_ALLOWLIST)


def looks_like_org(text: str) -> bool:
    return bool(re.search(
        r"\b(Limited|Ltd\.?|LLP|Pvt\.?\s?Ltd\.?|Private\s+Limited|Trust|HUF|Inc\.?|Corporation)\b",
        text, re.IGNORECASE))


class PIIRedactor:
    def __init__(self):
        self.fake = FakeIdentityProvider()
        self.registry = NameRegistry()
        self.detections = []  # flat log for the evaluation report
        self.free_text_detectors = [
            EmailDetector(), IPAddressDetector(), SSNDetector(),
            CreditCardDetector(), PhoneNumberDetector(), DateOfBirthDetector(),
            AddressDetector(), HonorificNameDetector(), CompanySuffixDetector(),
        ]

    # -- registry -----------------------------------------------------
    def learn(self, text: str, kind: str):
        """Record a name/org discovered via a structural anchor."""
        text = text.strip()
        if not text or text in ("[\u25cf]", "[●]"):
            return
        if kind == "PERSON":
            self.registry.add_person(text)
        elif kind == "ORG":
            if not is_allowlisted(text):
                self.registry.add_org(text)
        elif kind == "PERSON_OR_ORG":
            if looks_like_org(text):
                if not is_allowlisted(text):
                    self.registry.add_org(text)
            else:
                self.registry.add_person(text)

    # -- fake replacement dispatch -------------------------------------
    def _fake_for(self, label: str, original: str) -> str:
        return {
            "PERSON": self.fake.person_name,
            "EMAIL": self.fake.email,
            "COMPANY": self.fake.company,
            "ADDRESS": self.fake.address,
            "PHONE_NUMBER": self.fake.phone,
            "SSN": self.fake.ssn,
            "CREDIT_CARD": self.fake.credit_card,
            "IP_ADDRESS": self.fake.ip_address,
            "DATE_OF_BIRTH": self.fake.date_of_birth,
        }[label](original)

    def _log(self, label, original, fake, location):
        self.detections.append({
            "type": label, "original_len": len(original), "fake": fake,
            "location": location,
        })

    # -- whole-cell redaction (structural pass) ------------------------
    def redact_cell_full(self, text: str, kind: str, location: str) -> str:
        text = text.strip()
        if not text or text in ("[\u25cf]", "[●]", "N.A.", "NA", "-"):
            return text
        if kind == "PERSON":
            fake = self.fake.person_name(text)
            self._log("PERSON", text, fake, location)
            return fake
        if kind == "ORG":
            if is_allowlisted(text):
                return text
            fake = self.fake.company(text)
            self._log("COMPANY", text, fake, location)
            return fake
        if kind == "PERSON_OR_ORG":
            if looks_like_org(text):
                if is_allowlisted(text):
                    return text
                fake = self.fake.company(text)
                self._log("COMPANY", text, fake, location)
                return fake
            fake = self.fake.person_name(text)
            self._log("PERSON", text, fake, location)
            return fake
        if kind == "ADDRESS":
            fake = self.fake.address(text)
            self._log("ADDRESS", text, fake, location)
            return fake
        if kind == "EMAIL":
            # a cell may contain more than one email
            return self._sub_pattern(text, EmailDetector(), "EMAIL", location)
        if kind == "PHONE_TEXT":
            return self._sub_pattern(text, PhoneNumberDetector(), "PHONE_NUMBER", location)
        if kind == "EMAIL_PHONE":
            text = self._sub_pattern(text, EmailDetector(), "EMAIL", location)
            text = self._sub_pattern(text, PhoneNumberDetector(), "PHONE_NUMBER", location)
            return text
        return text

    def _sub_pattern(self, text, detector, label, location):
        spans = detector.find(text)
        for sp in sorted(spans, key=lambda s: -s.start):
            fake = self._fake_for(label, sp.text)
            self._log(label, sp.text, fake, location)
            text = text[:sp.start] + fake + text[sp.end:]
        return text

    # -- free text pass --------------------------------------------------
    def redact_free_text(self, text: str, location: str) -> str:
        if not text or not text.strip():
            return text

        spans = []
        for det in self.free_text_detectors:
            spans.extend(det.find(text))

        # Exact-match sweep for names/orgs discovered structurally elsewhere.
        for term, label in self.registry.all_terms():
            for m in re.finditer(re.escape(term), text, re.IGNORECASE):
                spans.append(Span(m.start(), m.end(), m.group(), label))

        if not spans:
            return text

        # Resolve overlaps: longer span wins; ties keep the first found.
        spans.sort(key=lambda s: (s.start, -(s.end - s.start)))
        chosen = []
        last_end = -1
        for sp in sorted(spans, key=lambda s: (-(s.end - s.start), s.start)):
            if any(not (sp.end <= c.start or sp.start >= c.end) for c in chosen):
                continue
            chosen.append(sp)
        chosen.sort(key=lambda s: -s.start)  # right-to-left for safe splicing

        for sp in chosen:
            label = sp.label
            if label == "COMPANY" and is_allowlisted(sp.text):
                continue
            fake = self._fake_for(label if label != "COMPANY" else "COMPANY", sp.text)
            self._log(label, sp.text, fake, location)
            text = text[:sp.start] + fake + text[sp.end:]
        return text


# ---------------------------------------------------------------------------
# docx structural walk helpers
# ---------------------------------------------------------------------------

def iter_body_blocks(document):
    """Yield ('p', Paragraph) / ('tbl', Table) in document order."""
    parent_elm = document.element.body
    for child in parent_elm.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            yield "p", Paragraph(child, document)
        elif tag == "tbl":
            yield "tbl", Table(child, document)


def classify_header(header_text: str):
    h = header_text.strip().lower()
    if h in COLUMN_SKIP:
        return None
    for key, kind in COLUMN_TYPE_MAP:
        if key in h:
            return kind
    return None


def row_is_structurally_reliable(row) -> bool:
    """
    Detects horizontally-merged cells. python-docx represents a merged
    cell by returning the *same* underlying XML element for every column
    it spans, so column-index-based classification silently corrupts once
    a row contains a merge (a disclaimer paragraph living in a full-width
    merged row would otherwise get relabelled as, say, a "Contact Person"
    name because it happens to sit under that column index).

    If every cell in the row maps to a distinct XML element, column-index
    classification is safe for this row; otherwise we fall back to the
    free-text pass for the whole row.
    """
    seen_ids = set()
    for c in row.cells:
        tc_id = id(c._tc)
        if tc_id in seen_ids:
            return False
        seen_ids.add(tc_id)
    return True


# -- PASS 1: learn names/orgs from structural anchors ------------------------

PROMOTERS_RE = re.compile(
    r"OUR\s+PROMOTERS\s*:\s*(.+?PRIVATE LIMITED|.+?LIMITED|.+?LLP)",
    re.IGNORECASE | re.DOTALL,
)
CONTACT_PERSON_RE = re.compile(
    r"Contact Person\s*:?\s*([A-Z][a-zA-Z.'-]+(?:\s+[A-Z][a-zA-Z.'-]+){0,3})"
)
IS_OUR_RE = re.compile(
    r"\b([A-Z][a-zA-Z.'-]+(?:\s+[A-Z][a-zA-Z.'-]+){1,3})\s+is our\b"
)


def learn_from_text(text: str, redactor: "PIIRedactor"):
    """Context-anchored name/org harvesting -- applied to both paragraph
    text and table-cell text, since either can hold the same patterns."""
    if not text:
        return
    m = PROMOTERS_RE.search(text)
    if m:
        body = m.group(1)
        parts = re.split(r",|\bAND\b", body)
        parts = [p.strip() for p in parts if p.strip()]
        for p in parts:
            if looks_like_org(p) or "TRUST" in p.upper() or "HUF" in p.upper():
                redactor.learn(p.title(), "ORG")
            else:
                redactor.learn(p.title(), "PERSON")
    for m in CONTACT_PERSON_RE.finditer(text):
        redactor.learn(m.group(1), "PERSON")
    for m in IS_OUR_RE.finditer(text):
        redactor.learn(m.group(1), "PERSON")


def learn_pass(document, redactor: PIIRedactor):
    for kind, block in iter_body_blocks(document):
        if kind == "tbl":
            if not block.rows:
                continue
            header_cells = [c.text for c in block.rows[0].cells]
            col_kinds = [classify_header(h) for h in header_cells]
            header_lower = {h.strip().lower() for h in header_cells}
            for row in block.rows:
                for cell in row.cells:
                    learn_from_text(cell.text, redactor)
            if not any(col_kinds):
                continue
            for row in block.rows[1:]:
                if not row_is_structurally_reliable(row):
                    continue  # merged row -- content doesn't line up with the header columns
                for cell, ck in zip(row.cells, col_kinds):
                    if ck in ("PERSON", "PERSON_OR_ORG"):
                        val = cell.text.strip()
                        if (val and val not in ("[●]",)
                                and val.lower() not in header_lower  # repeated header row
                                and not val.lower().startswith(("total", "sub-total"))):
                            redactor.learn(val, "PERSON_OR_ORG" if ck == "PERSON_OR_ORG" else "PERSON")
        else:
            learn_from_text(block.text, redactor)


# -- PASS 2: build the redacted output docx ----------------------------------

def copy_paragraph_redacted(src_para: Paragraph, out_doc, redactor: PIIRedactor, loc: str):
    text = src_para.text
    redacted = redactor.redact_free_text(text, loc) if text.strip() else text
    style_name = src_para.style.name if src_para.style else "Normal"
    # NOTE: python-docx's add_paragraph(text, style=X) appends the paragraph
    # to the document *before* resolving the style name, so a KeyError for
    # an unknown style (this blank output doc doesn't inherit the source's
    # style set) still leaves a half-added paragraph behind. Add the text
    # unstyled first, then apply the style separately so a failed lookup
    # can't create a duplicate paragraph.
    p = out_doc.add_paragraph(redacted)
    try:
        p.style = style_name
    except KeyError:
        pass
    return p


def copy_table_redacted(src_table: Table, out_doc, redactor: PIIRedactor, loc: str):
    n_rows = len(src_table.rows)
    n_cols = len(src_table.columns)
    if n_rows == 0 or n_cols == 0:
        return
    header_cells = [c.text for c in src_table.rows[0].cells]
    col_kinds = [classify_header(h) for h in header_cells]
    structural = any(col_kinds)

    out_table = out_doc.add_table(rows=0, cols=n_cols)
    try:
        out_table.style = "Light Grid Accent 1"
    except KeyError:
        pass

    for r_idx, row in enumerate(src_table.rows):
        out_row = out_table.add_row()
        row_reliable = row_is_structurally_reliable(row)
        seen_tc = {}  # id(source _tc) -> first output column index in this row (dedupe merges)
        for c_idx, cell in enumerate(row.cells):
            if c_idx >= n_cols:
                continue
            tc_id = id(cell._tc)
            if tc_id in seen_tc:
                # Same merged source cell as an earlier column -- mirror the
                # merge in the output instead of duplicating redacted text.
                try:
                    out_row.cells[seen_tc[tc_id]].merge(out_row.cells[c_idx])
                except Exception:
                    pass
                continue
            seen_tc[tc_id] = c_idx
            cell_loc = f"{loc}[r{r_idx}c{c_idx}]"
            raw = cell.text
            if structural and r_idx > 0 and row_reliable and col_kinds[c_idx]:
                redacted = redactor.redact_cell_full(raw, col_kinds[c_idx], cell_loc)
            else:
                redacted = redactor.redact_free_text(raw, cell_loc) if raw.strip() else raw
            out_row.cells[c_idx].text = redacted


def process_document(input_path: str, output_path: str):
    src = docx.Document(input_path)
    redactor = PIIRedactor()

    # Pass 1: learn real names/orgs from clearly-structured anchors.
    learn_pass(src, redactor)

    # Pass 2: rebuild the document, redacting as we go.
    out = docx.Document()
    for kind, block in iter_body_blocks(src):
        if kind == "p":
            copy_paragraph_redacted(block, out, redactor, "body")
        else:
            copy_table_redacted(block, out, redactor, "table")

    out.save(output_path)
    return redactor.detections


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("input", help="source .docx")
    ap.add_argument("output", help="redacted .docx to write")
    ap.add_argument("--log", help="optional path to write a JSON detection log")
    args = ap.parse_args()

    detections = process_document(args.input, args.output)

    by_type = {}
    for d in detections:
        by_type[d["type"]] = by_type.get(d["type"], 0) + 1
    print(f"Redacted {len(detections)} PII instances:")
    for t, c in sorted(by_type.items(), key=lambda x: -x[1]):
        print(f"  {t:15s} {c}")

    if args.log:
        Path(args.log).write_text(json.dumps(detections, indent=2))
        print(f"Detection log written to {args.log}")


if __name__ == "__main__":
    sys.exit(main())
